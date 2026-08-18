from __future__ import annotations

from pathlib import Path
import os
import zipfile

from docx import Document


ROOT = Path(__file__).resolve().parent
TARGET = ROOT / "FullPaperTemplateASIM2026 - KK.docx"
TEMP = ROOT / "FullPaperTemplateASIM2026 - KK.reflow_tmp.docx"
LONG_START = "Galapagos in Grasshopper was operated in minimization mode."
CAPTION_START = "Figure 3. Representative seven-occupant optimization:"


def main() -> None:
    doc = Document(TARGET)
    method_paragraph = None
    caption_paragraph = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(LONG_START):
            method_paragraph = paragraph
        elif text.startswith(CAPTION_START):
            caption_paragraph = paragraph
    if method_paragraph is None or caption_paragraph is None:
        raise ValueError("Required Figure 3 method elements were not found")

    method_element = method_paragraph._p
    method_element.getparent().remove(method_element)
    caption_paragraph._p.addnext(method_element)

    doc.save(TEMP)
    with zipfile.ZipFile(TEMP) as archive:
        bad = archive.testzip()
        if bad:
            raise ValueError(f"Corrupt ZIP member: {bad}")
    Document(TEMP)
    os.replace(TEMP, TARGET)
    print("moved_detailed_ga_method_after_figure3=true")


if __name__ == "__main__":
    main()
