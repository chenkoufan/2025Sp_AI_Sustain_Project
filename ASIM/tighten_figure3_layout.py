from __future__ import annotations

from pathlib import Path
import os
import re
import zipfile

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
TARGET = ROOT / "FullPaperTemplateASIM2026 - KK.docx"
TEMP = ROOT / "FullPaperTemplateASIM2026 - KK.layout_tmp.docx"
SCALE = 0.90


def scale_vml_style(style: str, scale: float) -> str:
    def replace(match: re.Match[str]) -> str:
        key, value = match.group(1), float(match.group(2))
        return f"{key}:{value * scale:.2f}pt"

    return re.sub(r"\b(width|height):([0-9.]+)pt", replace, style)


def main() -> None:
    doc = Document(TARGET)
    if len(doc.tables) < 3:
        raise ValueError("Representative optimization table was not found")
    table = doc.tables[2]
    shapes = table._tbl.xpath('.//*[local-name()="shape"]')
    if len(shapes) != 2:
        raise ValueError(f"Expected two Figure 3 VML images, found {len(shapes)}")
    for shape in shapes:
        style = shape.get("style") or ""
        shape.set("style", scale_vml_style(style, SCALE))

    # Let the row follow the scaled image height instead of retaining its older minimum.
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is not None:
        tr_pr.remove(tr_height)

    doc.save(TEMP)
    with zipfile.ZipFile(TEMP) as archive:
        bad = archive.testzip()
        if bad:
            raise ValueError(f"Corrupt ZIP member: {bad}")
    Document(TEMP)
    os.replace(TEMP, TARGET)
    print("scaled_figure3=0.90")


if __name__ == "__main__":
    main()
