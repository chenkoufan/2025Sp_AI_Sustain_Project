from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "ASIM2026_results_analysis_report.docx"
FIG = ROOT / "ASIM2026_figures"

BLUE = "087BB5"
LIGHT_BLUE = "DCEFF7"
LIGHT_GREY = "F2F4F5"
DARK = "222222"
ORANGE = "D95F02"


def set_run_font(run, size=10.5, bold=False, color=DARK, latin="Arial", east_asia="Microsoft YaHei"):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=DARK, size=8.8, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths=None, font_size=8.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, color="FFFFFF", size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr.cells[i], BLUE)
        if widths:
            hdr.cells[i].width = Cm(widths[i])
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        for i, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if i in (0, len(values) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(row.cells[i], value, size=font_size, align=align)
            if widths:
                row.cells[i].width = Cm(widths[i])
            if row_idx % 2 == 1:
                shade_cell(row.cells[i], LIGHT_GREY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    return table


def add_para(doc, text="", bold_lead=None, style=None, space_after=5, keep_with_next=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = keep_with_next
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=10.1)


def add_numbered(doc, items):
    for lead, body in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(lead)
        set_run_font(r1, size=10.1, bold=True, color=BLUE)
        r2 = p.add_run(body)
        set_run_font(r2, size=10.1)


def add_figure(doc, filename, caption, width_cm=17.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(FIG / filename), width=Cm(width_cm))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(7)
    cp.paragraph_format.keep_together = True
    r = cp.add_run(caption)
    set_run_font(r, size=8.5, color="444444")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8, color="777777")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=8, color="777777")


def style_document(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.65)
    sec.bottom_margin = Cm(1.55)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)
    sec.header_distance = Cm(0.7)
    sec.footer_distance = Cm(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color, before, after in (
        ("Title", 17, DARK, 0, 5),
        ("Subtitle", 11, "5A5A5A", 0, 12),
        ("Heading 1", 12.5, BLUE, 10, 5),
        ("Heading 2", 11.0, DARK, 7, 4),
    ):
        s = doc.styles[name]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.bold = name != "Subtitle"
        s.font.color.rgb = RGBColor.from_string(color)
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("ASIM 2026 · Results Analysis")
    set_run_font(r, size=8, color="777777")
    add_page_number(sec.footer.paragraphs[0])


def build():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(title.add_run("ASIM 2026 照明控制实验结果分析报告"), size=17, bold=True)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(subtitle.add_run("从初步现象到多因素交互结论"), size=11, color="5A5A5A")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    r = meta.add_run("数据范围：3 种天空条件 × 15 个占用人数 × 5 种座位布局 = 225 个配对场景｜指标：照明功率（W）")
    set_run_font(r, size=8.8, color="666666")

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.autofit = False
    cell = callout.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("核心结论｜")
    set_run_font(r1, size=10.3, bold=True, color=BLUE)
    r2 = p.add_run(
        "混合顶灯—桌面灯控制相较纯顶灯分区 PIR 显著降低功率；GA 的额外价值则取决于运行情境："
        "Clear 下规则控制已接近功率下限，Overcast 和 Night 下 GA 分别进一步降低 27.02 W 和 78.68 W。"
        "优化价值由日光可用性、占用密度、座位空间分布和控制架构共同决定。"
    )
    set_run_font(r2, size=10.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    doc.add_heading("1 分析基础：先保证比较公平", level=1)
    add_para(doc,
        "数据覆盖 Clear、Overcast、Night 三种天空条件；每种条件包含 1–15 人、每个人数五种座位布局，共 75 个场景。"
        "所有策略均在相同天气、相同人数、相同座位组合下配对比较。五种布局属于预设设计样本，并非总体随机重复，"
        "因此报告均值、差值、范围和完整胜负次数，不作显著性检验。功率单位为 W，不等同于时段能耗。"
    )
    add_table(doc,
        ["天空", "分区 PIR\n纯顶灯", "规则混合", "GA 混合", "GA 相对 PIR", "GA 相对规则"],
        [
            ["Clear", "75.50 W", "16.34 W", "17.44 W", "−58.06 W\n(−76.9%)", "+1.10 W\n(+6.7%)"],
            ["Overcast", "244.46 W", "120.03 W", "93.01 W", "−151.45 W\n(−62.0%)", "−27.02 W\n(−22.5%)"],
            ["Night", "443.52 W", "314.78 W", "236.10 W", "−207.42 W\n(−46.8%)", "−78.68 W\n(−25.0%)"],
        ],
        widths=[2.0, 2.8, 2.5, 2.5, 3.3, 3.3],
        font_size=8.4,
    )

    doc.add_page_break()
    doc.add_heading("2 初步结论：天空条件决定负荷水平，人数只解释部分变化", level=1)
    add_figure(
        doc,
        "results_power_by_occupancy.png",
        "图1｜三种天空条件下功率随占用人数的变化。折线为五种匹配座位布局的均值；阴影为最小值至最大值，并非置信区间。",
        width_cm=17.0,
    )
    add_para(doc,
        "首先，日光越弱，三种策略的功率均越高，天空条件构成照明需求的首要边界。其次，功率总体随人数增加，但并不严格单调；"
        "分区 PIR 尤其明显，因为相同人数落在不同区域会触发不同顶灯档位。再次，两种混合控制在三种天气下均大幅低于纯顶灯 PIR，"
        "但 Clear 条件中规则混合与 GA 几乎重合，提示主要收益来自控制架构，而不是进一步搜索。"
    )
    add_para(doc,
        "阴影带进一步说明：人数相同不代表照明需求相同。座位靠窗、靠内区或跨越多个控制区，会改变可利用日光和顶灯触发范围。"
        "因此，总体均值只能形成第一层结论，策略比较还必须控制座位布局。",
        bold_lead="阴影带进一步说明：",
    )

    doc.add_page_break()
    doc.add_heading("3 条件化结论：GA 的优势随天空条件和密度发生转折", level=1)
    add_figure(
        doc,
        "deep_paired_strategy_evidence.png",
        "图2｜严格配对后的策略证据。(a,b) 每格为五种布局的平均配对节省，正值表示 GA 功率更低；(c) 汇总每种比较的 75 个场景。",
        width_cm=17.0,
    )
    add_para(doc,
        "相对分区 PIR，GA 在 225 个场景中有 217 个功率更低、5 个相同、3 个更高，显示混合优化总体上稳定优于纯顶灯控制。"
        "相对更强的规则混合基线，结论则具有明确条件性。"
    )
    add_bullets(doc, [
        "Clear：GA 仅在 7/75 个场景更低，28 个相同，40 个更高；平均多用 1.10 W，任何人数下均未形成正的平均节省。",
        "Overcast：GA 在 60/75 个场景更低；4–15 人的 60 个场景中，59 个更低、1 个相同、0 个更高。",
        "Night：1 人时五种布局全部持平；2–15 人中，GA 有 69 个更低、1 个相同、0 个更高。",
    ])
    add_para(doc,
        "不能把所有密度混合后只给一个总平均。更准确的结论是：GA 的增益在日光不足且存在多工位协调需求时出现；日光充足时，简单规则已足够。",
        bold_lead="不能把所有密度混合后只给一个总平均。",
    )

    doc.add_page_break()
    doc.add_heading("4 机理结论：优势来自重新分配，而非等比例调暗", level=1)
    add_figure(
        doc,
        "deep_mechanism_and_layout.png",
        "图3｜优化机理与空间敏感性。(a) GA 减去规则混合，负值表示 GA 更低；(b) 同一天气和人数下五种布局的平均极差，柱顶百分数为布局内变化占总变化的比例。",
        width_cm=17.0,
    )
    add_table(doc,
        ["天空", "顶灯变化", "桌面灯变化", "总功率变化", "机理"],
        [
            ["Clear", "0.00 W", "+1.10 W", "+1.10 W", "顶灯已关闭，优化空间很小"],
            ["Overcast", "−25.70 W", "−1.32 W", "−27.02 W", "主要降低背景顶灯"],
            ["Night", "−84.30 W", "+5.62 W", "−78.68 W", "少量任务光替代大量背景顶灯"],
        ],
        widths=[1.8, 2.6, 2.8, 2.8, 6.3],
        font_size=8.4,
    )
    add_para(doc,
        "Night 最能体现协调机制：桌面灯平均增加 5.62 W，却换来顶灯减少 84.30 W，净降 78.68 W。"
        "若只报告总功率，会遗漏“局部补光替代全局照明”这一关键发现。"
    )

    doc.add_page_break()
    doc.add_heading("5 高维结论：四个因素共同决定最优策略", level=1)
    add_numbered(doc, [
        ("天空条件定义可优化空间。", "Clear 时日光已承担主要照明；Overcast 和 Night 时人工照明成为主要可控负荷。"),
        ("占用密度决定协调复杂度。", "Overcast 的 GA 优势从 4 人开始稳定；Night 在 2 人后几乎全面出现。中密度 Night 的平均节省最大，为 100.24 W（28.8%）。"),
        ("座位布局决定同密度下的实际需求。", "Clear 条件中，布局内变化占总变化的 26.0%–48.8%，说明有日光时“坐在哪里”不能被“有几个人”替代。"),
        ("控制架构决定能否利用空间差异。", "PIR 只能按区开启顶灯；规则混合先获得大部分架构收益；GA 在弱日光和多人分布下继续协调顶灯与桌面灯。"),
    ])
    add_table(doc,
        ["条件", "主要收益来源", "GA 相对规则混合", "结论强度"],
        [
            ["Clear", "混合照明架构", "无额外节能优势", "近似等效；不宣称 GA 优越"],
            ["Overcast", "架构 + 档位协调", "4 人以上跨布局稳定受益", "条件性优势明确"],
            ["Night", "局部任务光替代背景顶灯", "2 人以上几乎全面受益", "最强且最稳定"],
        ],
        widths=[2.0, 4.6, 5.0, 5.0],
        font_size=8.6,
    )
    add_para(doc,
        "高维控制逻辑不是为所有工况选定一个固定赢家，而是形成分情境策略：日光充足时采用规则混合控制；日光不足、人数增加或座位跨区分散时启用 GA。"
        "若未来部署需要控制计算成本，可进一步形成以天空条件和占用状态为门控变量的分层控制器。"
    )

    doc.add_heading("6 论文级最终结论", level=1)
    final_p = doc.add_table(rows=1, cols=1)
    final_p.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = final_p.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "混合环境照明与任务照明相较分区 PIR 纯顶灯控制，在三种天空条件下均显著降低照明功率。"
        "然而，GA 的附加收益具有明确情境依赖性：Clear 下规则混合已接近最低功率；Overcast 和 Night 下，GA 通过减少背景顶灯并将光输出定向分配至占用工位，"
        "分别取得 22.5% 和 25.0% 的进一步平均节省。座位布局在同一人数下仍造成明显功率差异，表明占用人数不足以完整描述照明需求。"
        "本研究因而界定了日光可用性、占用密度、空间分布与控制架构共同作用下，优化控制真正有价值的运行区间。"
    )
    set_run_font(r, size=10.1)

    doc.add_heading("7 结论边界与下一步核查", level=1)
    add_bullets(doc, [
        "当前表格未包含每个最终解的照度与均匀度输出；功率比较已经验证，但仍需回查各策略是否满足完全相同的视觉约束。",
        "回查三个 GA 高于零功率 PIR 的场景：Clear 两人（13、18），Clear 三人（13、17、22），Overcast 一人（23）。",
        "正式论文宜补充重复优化或收敛检查，并将“全局最优”表述为“在给定搜索设置下获得的优化解”。",
    ])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("一句话结论｜")
    set_run_font(r1, size=10.5, bold=True, color=ORANGE)
    r2 = p.add_run("混合照明决定主要节能幅度，GA 决定弱日光、多工位和空间分散条件下还能进一步节省多少。")
    set_run_font(r2, size=10.5, bold=True)

    core = doc.core_properties
    core.title = "ASIM 2026 照明控制实验结果分析报告"
    core.subject = "从初步现象到多因素交互结论"
    core.author = "ASIM 2026 study team"
    core.keywords = "lighting control, occupancy, daylight, genetic algorithm, ASIM 2026"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
