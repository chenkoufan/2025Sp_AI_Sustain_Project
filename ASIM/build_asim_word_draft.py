from __future__ import annotations

import argparse
from pathlib import Path


ROOT = Path(__file__).resolve().parent
SIM = ROOT / "ASIM-simulation"
FIGDIR = ROOT / "ASIM2026_figures"
TARGET = ROOT / "FullPaperTemplateASIM2026 - KK.docx"
BACKUP = ROOT / "FullPaperTemplateASIM2026 - KK.original.docx"


def build_figures() -> None:
    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    from PIL import Image

    FIGDIR.mkdir(exist_ok=True)

    # Figure 1: reuse the case-model images already embedded in the local report.
    media = ROOT / ".tmp_bps_extract" / "word" / "media"
    building = Image.open(FIGDIR / "building.png").convert("RGB")
    plan = Image.open(media / "image2.png").convert("RGB")
    perspective = Image.open(media / "image5.png").convert("RGB")
    fig, axes = plt.subplots(1, 3, figsize=(7.0, 3.15), layout="constrained")
    for ax, im, label, title in [
        (axes[0], building, "(a)", "Simulated room in host building"),
        (axes[1], plan, "(b)", "Workstation and lighting layout"),
        (axes[2], perspective, "(c)", "Simulation model"),
    ]:
        ax.imshow(im)
        ax.axis("off")
        ax.set_title(f"{label} {title}", fontsize=9, loc="left", pad=3)
    fig.savefig(FIGDIR / "figure1_case_model.png", dpi=300, facecolor="white")
    plt.close(fig)

    colors = {"Clear": "#0072B2", "Overcast": "#E69F00", "Night": "#333333"}
    markers = {"Clear": "o", "Overcast": "s", "Night": "^"}
    frames = {}
    for condition in ["Clear", "Overcast", "Night"]:
        frames[condition] = pd.read_excel(SIM / f"dataset_{condition}.xlsx")

    fig, axes = plt.subplots(1, 2, figsize=(7.0, 3.05), layout="constrained")
    ax = axes[0]
    for condition, df in frames.items():
        grouped = df.groupby("occupancy_count")["power"]
        x = np.arange(1, 16)
        mean = grouped.mean().reindex(x).to_numpy()
        lo = grouped.min().reindex(x).to_numpy()
        hi = grouped.max().reindex(x).to_numpy()
        ax.fill_between(x, lo, hi, color=colors[condition], alpha=0.14, linewidth=0)
        ax.plot(
            x,
            mean,
            label=condition,
            color=colors[condition],
            marker=markers[condition],
            markersize=3.5,
            linewidth=1.4,
        )
    ax.set(xlabel="Occupancy count", ylabel="Optimized lighting power (W)", xlim=(1, 15))
    ax.set_xticks([1, 3, 5, 7, 9, 11, 13, 15])
    ax.set_ylim(bottom=0)
    ax.grid(axis="y", color="#D9D9D9", linewidth=0.6)
    ax.legend(frameon=False, fontsize=8, loc="upper left")
    ax.set_title("(a) Mean and range across five seating patterns", fontsize=9, loc="left")

    ax = axes[1]
    conditions = ["Clear", "Overcast", "Night"]
    ceiling_means = []
    task_means = []
    for condition in conditions:
        df = frames[condition]
        ceiling = [c for c in df.columns if c.startswith("ceiling_")]
        task = [c for c in df.columns if c.startswith("task_")]
        ceiling_means.append((13.2 * df[ceiling].sum(axis=1)).mean())
        task_means.append((1.5 * df[task].sum(axis=1)).mean())
    x = np.arange(len(conditions))
    ax.bar(x, ceiling_means, color="#999999", edgecolor="black", linewidth=0.5, label="Ceiling")
    ax.bar(
        x,
        task_means,
        bottom=ceiling_means,
        color="#56B4E9",
        edgecolor="black",
        linewidth=0.5,
        hatch="//",
        label="Task",
    )
    ax.set_xticks(x, conditions)
    ax.set(ylabel="Mean optimized power (W)", ylim=(0, 270))
    ax.grid(axis="y", color="#D9D9D9", linewidth=0.6)
    ax.legend(frameon=False, fontsize=8, loc="upper left")
    ax.set_title("(b) Mean power allocation", fontsize=9, loc="left")
    for i, (c, t) in enumerate(zip(ceiling_means, task_means)):
        total = c + t
        ax.text(i, total + 5, f"{total:.1f}", ha="center", va="bottom", fontsize=8)
    fig.savefig(FIGDIR / "figure2_results.png", dpi=300, facecolor="white")
    plt.close(fig)


def build_document() -> None:
    import shutil

    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    if not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)

    doc = Document(TARGET)
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = False
    for area in [
        section.header,
        section.first_page_header,
        section.even_page_header,
        section.footer,
        section.first_page_footer,
        section.even_page_footer,
    ]:
        area.is_linked_to_previous = False
        for p in area.paragraphs:
            p.clear()
            if p._p.pPr is not None:
                p._p.remove(p._p.pPr)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def set_run(run, *, size=12, bold=False, italic=False, color=None, highlight=False):
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
        if highlight:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "FFF2CC")
            run._element.get_or_add_rPr().append(shd)
        return run

    def paragraph(text="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, before=0, keep=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.keep_with_next = keep
        if text:
            set_run(p.add_run(text))
        return p

    def major(text):
        p = paragraph(align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3, keep=True)
        set_run(p.add_run(text.upper()), bold=True)
        return p

    def subheading(text):
        p = paragraph(align=WD_ALIGN_PARAGRAPH.LEFT, before=3, after=0, keep=True)
        set_run(p.add_run(text), bold=True)
        return p

    def add_caption(text):
        p = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=6, keep=False)
        set_run(p.add_run(text), bold=True, italic=True)
        return p

    def shade_cell(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    # Front matter
    p = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, after=12, keep=True)
    set_run(
        p.add_run("BALANCING BACKGROUND AND TASK LIGHTING UNDER DYNAMIC OCCUPANCY: A SIMULATION-BASED OPTIMIZATION STUDY"),
        size=14,
        bold=True,
    )
    p = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, after=6, keep=True)
    set_run(p.add_run("Kefan Chen¹, Guanli Feng², Yu Qian Ang¹"))
    p = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, after=0, keep=True)
    set_run(p.add_run("¹Department of Built Environment, National University of Singapore, Singapore"))
    p = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, after=6, keep=True)
    set_run(p.add_run("²Department of Architecture, National University of Singapore, Singapore"))
    p = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, after=12, keep=True)
    set_run(p.add_run("Corresponding author: [email to be added]"), italic=True, highlight=True)

    major("Abstract")
    paragraph(
        "Traditional lighting design commonly assumes full occupancy, whereas workplace occupancy varies in both density and spatial distribution. This mismatch can cause unnecessary illumination in underused spaces. Background lighting provides room-scale ambient illumination, while task lighting serves individual workstations; coordinating the two may therefore improve the response to changing occupancy. This study proposes a Balanced Light Control Framework to investigate the minimum lighting power achievable while maintaining visual-comfort requirements. A three-dimensional model of a Singapore office laboratory was evaluated under three representative lighting conditions and 15 occupancy levels, from one to 15 occupants. Five random seating distributions were generated at each level and applied consistently across the three conditions, producing 225 optimization scenarios. Three ceiling-light zones and 24 individual task lights were treated as decision variables. A penalty-based objective function combined lighting power with penalties for failing illuminance and uniformity thresholds from SS 531:2006. A genetic algorithm implemented in Galapagos identified a minimum-power feasible configuration for each scenario. Mean optimized power was 17.62 W under clear conditions, 93.01 W under overcast conditions, and 236.10 W at night. Ceiling lighting contributed 1.0%, 65.7%, and 82.4% of optimized power, respectively. The results reveal a daylight-dependent transition from task-light-dominant to background-light-dominant operation. They also show that seating distribution, not occupant count alone, can materially affect lighting demand. The framework provides a case-specific performance benchmark for assessing occupancy-centric lighting strategies; comparative savings will be quantified after consistent control baselines are established."
    )
    p = paragraph(after=12)
    set_run(p.add_run("KEYWORDS"), bold=True)
    set_run(p.add_run("  Occupancy-centric lighting, task lighting, lighting optimization, visual comfort, daylight availability"))

    major("1. Introduction")
    paragraph(
        "Lighting systems in offices are often designed and commissioned for nominal or full occupancy, although actual use varies over time and across workstations. Controls based only on schedules or aggregate occupant counts may therefore illuminate areas that are unoccupied or already receive sufficient daylight. Room-level occupancy datasets further demonstrate that office use is temporally variable, motivating controls that respond to actual space use (Tekler et al. 2022)."
    )
    paragraph(
        "Centralized ceiling lighting and personalized task lighting operate at different spatial scales. Ceiling luminaires provide general illumination and uniformity, whereas task lights deliver illuminance close to individual workplanes. Personalized lighting can improve occupants' ability to adjust their local environment (De Korte et al. 2015), and prior work has examined the integration of task lighting with daylight and electric lighting (Papinutto et al. 2021). However, the minimum-power division between background and task lighting remains sensitive to daylight, occupancy density, and the locations of occupied seats. These factors are rarely examined together in a controlled scenario matrix."
    )
    paragraph(
        "This study asks: how should ceiling and task-light levels be coordinated under changing daylight and occupancy patterns to minimize lighting power while satisfying visual-comfort constraints? A Balanced Light Control Framework is evaluated in a simulated open-plan office. The contribution is threefold: (1) a joint representation of three ceiling-light zones and 24 task lights; (2) a penalty-based optimization formulation incorporating illuminance and uniformity; and (3) a matched experiment comprising 225 lighting–occupancy–seating scenarios. The emphasis is a theoretical, case-specific performance benchmark rather than a deployed real-time controller."
    )

    major("2. Methodology")
    subheading("2.1 Case-study model and lighting systems")
    paragraph(
        "A three-dimensional model was constructed for a single open-plan office laboratory located within a larger building in Singapore, using the existing architectural geometry. Figure 1a identifies the simulated room within the host-building model; only this room, rather than the whole building, was included in the lighting simulations. The room contains 24 fixed workstations, each equipped with an independently adjustable task light. The background system is divided into three ceiling-light zones: seats 0–7, 8–15, and 16–23. ClimateStudio was used to calculate illuminance and uniformity for the occupied task areas and their surroundings. Material reflectance, luminaire photometry, model orientation, and calculation-grid settings should be reported in the camera-ready version."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run()
    r.add_picture(str(FIGDIR / "figure1_case_model.png"), width=Cm(15.5))
    add_caption("Figure 1. Case-study context and model: (a) location of the simulated room within the host building; (b) workstation and lighting layout; (c) room-scale simulation model and calculation grid.")

    subheading("2.2 Scenario matrix")
    paragraph(
        "The experiment considered three representative conditions labelled Clear, Overcast, and Night. Daylight simulations were referenced to Changi International Airport, Singapore. The Overcast case used the ClimateStudio sky identifier ‘Sky:[Changi.Intl.AP,SG,SGP CIE_Overcast 3.21.13.00]’, corresponding to a CIE Overcast sky at 13:00 on 21 March. The exact software identifiers for the Clear and Night cases remain to be recorded. Occupancy count ranged from one to 15. At each count, five unique seating distributions were randomly selected from the 24 workstations. The same 75 occupancy–seating patterns were evaluated under all three conditions, enabling matched comparisons and yielding 225 optimization scenarios. The current dataset contains no missing or duplicated scenarios."
    )
    table_caption = add_caption("Table 1. Experimental scenario matrix.")
    table_caption.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    headers = ["Factor", "Levels", "Cases"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(cell.paragraphs[0].add_run(h), bold=True)
    rows = [
        ("Lighting condition", "Clear, Overcast, Night", "3"),
        ("Occupancy count", "1–15 occupants", "15"),
        ("Seating distribution", "Five patterns per count", "5"),
        ("Total", "Matched factorial matrix", "225"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(cells[i].paragraphs[0].add_run(value))
    for row in table.rows[:-1]:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_with_next = True

    subheading("2.3 Decision variables and lighting power")
    paragraph(
        "The control vector comprises discrete output levels for the three ceiling zones and 24 task lights. Task lights at unoccupied seats were constrained to remain off. The exported results indicate that each task-light level adds 1.5 W and each ceiling-zone level adds 13.2 W. Total lighting power was therefore calculated as:"
    )
    p = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=3)
    set_run(p.add_run("P = 13.2 Σⱼ₌₁³ Lc,j + 1.5 Σᵢ₌₁²⁴ Lt,i ."), italic=True)
    paragraph(
        "This expression reproduced the reported power for all 225 exported solutions without error. Because the simulations represent discrete conditions without an assigned duration, the analysis reports electrical power (W), not energy consumption (kWh)."
    )

    subheading("2.4 Visual-comfort constraints and objective function")
    paragraph(
        "Visual-comfort requirements were based on SS 531:2006. Occupied task areas were required to achieve at least 500 lux with uniformity Emin/Eavg ≥ 0.70; surrounding areas required at least 300 lux with uniformity ≥ 0.50. Violations were normalized and added to lighting power through a quadratic penalty function:"
    )
    p = paragraph(align=WD_ALIGN_PARAGRAPH.CENTER, before=3, after=3)
    set_run(p.add_run("Fitness = P + λE1(vE1)² + λE0(vE0)² + λU1(vU1)² + λU2(vU2)²."), italic=True)
    paragraph(
        "Here, v represents the normalized violation of each illuminance or uniformity threshold and λ is its penalty coefficient. The penalties prioritize feasible lighting conditions before minimizing power. The final manuscript should report the numerical penalty coefficients and any sensitivity test used to confirm that infeasible low-power solutions were excluded."
    )

    subheading("2.5 Genetic-algorithm optimization")
    paragraph(
        "Galapagos in Grasshopper was operated in minimization mode. For each scenario, the algorithm evolved combinations of the 27 discrete lighting variables. ClimateStudio and the associated calculation scripts evaluated task-area and surrounding-area illuminance, uniformity, and total power at each candidate solution. The retained output was the lowest-fitness configuration satisfying the encoded constraints. Population size, generation limit, convergence criterion, random seed, number of repeated runs, and software versions remain to be added for reproducibility."
    )
    paragraph(
        "Figure 2 documents a representative seven-occupant case, linking the complete Galapagos optimization interface to its associated seating distribution and simulated illuminance field."
    )
    gal_table = doc.add_table(rows=1, cols=2)
    gal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    gal_table.autofit = False
    gal_row = gal_table.rows[0]
    gal_tr_pr = gal_row._tr.get_or_add_trPr()
    gal_cant_split = OxmlElement("w:cantSplit")
    gal_tr_pr.append(gal_cant_split)
    gal_images = [
        (FIGDIR / "Screenshot 2026-08-10 145801.png", Cm(5.0)),
        (FIGDIR / "Galapagos-plan.png", Cm(2.8)),
    ]
    for cell, (image_path, width) in zip(gal_row.cells, gal_images):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_image = cell.paragraphs[0]
        p_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_image.paragraph_format.space_after = Pt(0)
        p_image.paragraph_format.keep_with_next = True
        p_image.add_run().add_picture(str(image_path), width=width)
    add_caption(
        "Figure 2. Representative seven-occupant optimization: (a) complete Galapagos interface; (b) corresponding occupant distribution and simulated illuminance field."
    )

    major("3. Results")
    subheading("3.1 Optimized power across lighting conditions")
    paragraph(
        "Optimized power differed sharply across the three conditions (Figure 3a). Averaged across the 75 occupancy–seating scenarios, it was 17.62 W for Clear, 93.01 W for Overcast, and 236.10 W for Night. The respective ranges were 0–46.2 W, 0–135.0 W, and 110.1–303.9 W. Power generally increased with occupancy, but it was not strictly monotonic because each occupancy level used a different set of spatial seating patterns."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run()
    r.add_picture(str(FIGDIR / "figure2_results.png"), width=Cm(14.2))
    add_caption("Figure 3. Optimized lighting power: (a) mean and full range across five seating distributions at each occupancy count; (b) mean allocation to ceiling and task lighting across all 75 scenarios per condition.")

    subheading("3.2 Shift between task and background lighting")
    paragraph(
        "The optimized division of power changed with daylight availability (Figure 3b). Ceiling lighting contributed only 1.0% of optimized power for Clear, compared with 65.7% for Overcast and 82.4% for Night. All ceiling zones were off in 74 of 75 Clear scenarios, five Overcast scenarios, and no Night scenarios. In contrast, the Night solutions used both lighting layers in every case. These results show a transition from almost entirely localized lighting under Clear conditions to background-dominant operation at Night."
    )

    subheading("3.3 Sensitivity to seating distribution")
    paragraph(
        "Solutions with the same number of occupants could require substantially different power. Under Overcast conditions at an occupancy count of three, the five patterns ranged from 9 to 84 W. At an occupancy count of two, Night solutions ranged from 113.1 to 196.8 W. Such within-count variation indicates that occupant count alone does not determine lighting demand; the relationship between occupied seats, daylight exposure, and the three ceiling zones also matters. Clear and Overcast cases with zero power were retained as valid zeros rather than missing data, indicating that the simulated daylight alone satisfied the encoded requirements at those seat locations."
    )

    major("4. Discussion")
    paragraph(
        "The principal finding is that background and task lighting have condition-dependent roles rather than a fixed optimal ratio. When daylight is abundant, localized task lighting can address residual workstation needs with little or no ceiling-light contribution. As daylight decreases, ceiling lighting becomes increasingly important for both ambient illumination and uniformity. At Night, the ceiling system supplied most of the optimized power, while task lights fine-tuned local requirements."
    )
    paragraph(
        "The variation among seating patterns has a direct control implication. A strategy based only on total occupancy would assign identical settings to spatially different cases, even though the optimized power can differ substantially. A practical occupancy-centric controller may therefore benefit from zonal or seat-level location information. The optimization results provide a theoretical reference against which simpler implementable controls can be evaluated. They do not, by themselves, demonstrate savings relative to PIR-based, schedule-based, ceiling-only, or rule-based combined controls."
    )
    paragraph(
        "A consistent baseline comparison is the most important remaining analysis. At minimum, future results should compare the optimized framework with: (1) ceiling-light-only control satisfying the same comfort constraints; and (2) a predefined combined strategy using fixed ceiling and task-light levels. A PIR-style zonal baseline would further respond to the concern that conventional occupancy sensing already reduces lighting in unoccupied zones. All strategies must use the same 225 scenarios and comfort criteria; otherwise percentage savings would not be directly comparable."
    )
    paragraph(
        "Several limitations bound interpretation. The study represents one office geometry and only three selected lighting conditions, rather than annual daylight variation. Occupancy was limited to 15 of 24 workstations, with five random distributions per count. The comfort model considered simulated illuminance and uniformity but not individual preference, glare, sensor uncertainty, control delay, or hardware constraints. Genetic algorithms can also return local or run-dependent solutions; repeated-run convergence evidence is needed before interpreting the outputs as a strict global performance ceiling."
    )

    major("5. Conclusion and Implications")
    paragraph(
        "This study formulated and evaluated a Balanced Light Control Framework that jointly optimizes three ceiling-light zones and 24 workstation task lights. Across 225 matched scenarios, mean optimized power increased from 17.62 W under Clear conditions to 93.01 W under Overcast conditions and 236.10 W at Night. The ceiling-light share simultaneously increased from 1.0% to 65.7% and 82.4%, demonstrating a daylight-dependent shift in the energy-minimizing lighting configuration. Considerable variation among seating patterns at the same occupancy count further shows the value of spatial occupancy information. The results establish a case-specific optimization benchmark; claims of relative energy savings should be reserved until comparable conventional-control baselines and reproducibility checks are completed."
    )

    major("Acknowledgements")
    p = paragraph()
    set_run(p.add_run("[Funding, technical assistance, and acknowledgements to be confirmed.]"), highlight=True)

    major("References")
    refs = [
        "De Korte EM, Spiekman M, Hoes-van Oeffelen L, Van Der Zande B, Vissenberg G, Huiskes G, and Kuijt-Evers LFM. 2015. Personal environmental control: Effects of pre-set conditions for heating and lighting on personal settings, task performance and comfort experience. Building and Environment 86: 166–176. https://doi.org/10.1016/j.buildenv.2015.01.002.",
        "Papinutto M, Colombo M, Golsouzidou M, Reutter K, Lalanne D, and Nembrini J. 2021. Towards the integration of personal task-lighting in an optimised balance between electric lighting and daylighting: A user-centred study of emotion, visual comfort, interaction and form-factor of task lights. Journal of Physics: Conference Series 2042: 012115. https://doi.org/10.1088/1742-6596/2042/1/012115.",
        "Singapore Standards Council. 2006. SS 531:2006 Code of Practice for Lighting of Work Places, Part 1: Indoor Lighting. Singapore: SPRING Singapore. [Edition and bibliographic details to verify.]",
        "Tekler ZD, Ono E, Peng Y, Zhan S, Lasternas B, and Chong A. 2022. ROBOD, room-level occupancy and building operation dataset. Building Simulation 15: 2127–2137. https://doi.org/10.1007/s12273-022-0925-9.",
    ]
    for ref in refs:
        p = paragraph(ref, after=0)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)

    # Enforce Times New Roman throughout, including tables and inherited runs.
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            if run.font.size is None:
                run.font.size = Pt(12)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        run.font.size = Pt(12)

    doc.core_properties.title = "Balancing Background and Task Lighting under Dynamic Occupancy"
    doc.core_properties.subject = "Working draft for ASIM 2026"
    doc.save(TARGET)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("mode", choices=["figures", "document"])
    args = parser.parse_args()
    if args.mode == "figures":
        build_figures()
    else:
        build_document()
