from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "FullPaperTemplateASIM2026 - KK - back.docx"
OUTPUT = ROOT / "FullPaperTemplateASIM2026 - KK.method_updated_v2.docx"
DIAGRAM = ROOT / "ASIM2026_figures" / "method_diagram.jpg"


OVERVIEW_HEADING = "Optimization workflow"
OVERVIEW_1 = (
    "The complete workflow is summarized in Figure 1. Environmental inputs comprised the "
    "three-dimensional office model, one of three sky conditions, and a randomly generated "
    "occupancy distribution containing 1–15 occupants (Figure 1(a)). These inputs were paired "
    "with a 27-variable lighting-control vector comprising the discrete output levels of three "
    "ceiling-light zones and 24 workstation task lights (Figure 1(b))."
)
OVERVIEW_2 = (
    "For each candidate vector, ClimateStudio calculated task-area and surrounding-area "
    "illuminance and uniformity, while the ceiling- and task-light levels were converted to total "
    "lighting power; these quantities formed the visual-comfort constraints and power objective "
    "(Figure 1(c)). The penalty-based genetic algorithm iteratively updated the lighting vector and "
    "repeated the simulation until the lowest-fitness feasible solution was retained (Figure 1(d)). "
    "This solution defined the balanced ceiling- and task-light configuration for the scenario "
    "(Figure 1(e)). Its power demand was then compared on a matched-case basis with the idealized "
    "zonal PIR ceiling-only and occupancy-triggered rule-based hybrid baselines (Figure 1(f))."
)
CAPTION_LABEL = "Figure 1. "
CAPTION_TEXT = (
    "Workflow of the simulation-based lighting-control study: (a) environmental inputs; "
    "(b) ceiling- and task-light decision variables; (c) lighting-power and visual-comfort "
    "objectives; (d) genetic-algorithm optimization; (e) optimized balanced control; and "
    "(f) matched comparison with the two rule-based baselines."
)


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def clear_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def copy_run_format(source_run, target_run) -> None:
    if source_run is not None and source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def set_plain_text(paragraph: Paragraph, text: str, run_template=None) -> None:
    if run_template is None and paragraph.runs:
        run_template = paragraph.runs[0]
    clear_content(paragraph)
    run = paragraph.add_run(text)
    copy_run_format(run_template, run)


def new_paragraph_after(anchor: Paragraph, template: Paragraph, text: str = "") -> Paragraph:
    element = deepcopy(template._p)
    clear = Paragraph(element, anchor._parent)
    clear_content(clear)
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    run = paragraph.add_run(text)
    copy_run_format(template.runs[0] if template.runs else None, run)
    return paragraph


def renumber_existing_figures(doc: Document) -> None:
    # Descending replacement prevents a newly assigned number from being shifted again.
    for old, new in ((5, 6), (4, 5), (3, 4), (2, 3), (1, 2)):
        needle = f"Figure {old}"
        replacement = f"Figure {new}"
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if needle in run.text:
                    run.text = run.text.replace(needle, replacement)


def add_method_figure(anchor: Paragraph, image_template: Paragraph) -> Paragraph:
    element = deepcopy(image_template._p)
    paragraph = Paragraph(element, anchor._parent)
    clear_content(paragraph)
    anchor._p.addnext(element)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run()
    shape = run.add_picture(str(DIAGRAM), width=Inches(6.65))
    shape._inline.docPr.set(
        "descr",
        "Workflow from environmental and occupancy inputs through lighting simulation, genetic-"
        "algorithm optimization, balanced ceiling and task-light control, and matched baseline comparison.",
    )
    shape._inline.docPr.set("title", "Simulation-based lighting-control workflow")
    return paragraph


def add_caption(anchor: Paragraph, template: Paragraph) -> Paragraph:
    paragraph = new_paragraph_after(anchor, template, "")
    clear_content(paragraph)
    label_run = paragraph.add_run(CAPTION_LABEL)
    label_run.bold = True
    label_run.italic = True
    text_run = paragraph.add_run(CAPTION_TEXT)
    text_run.italic = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    return paragraph


def main() -> None:
    if not SOURCE.exists() or not DIAGRAM.exists():
        raise FileNotFoundError("The source manuscript or method diagram is missing")

    doc = Document(SOURCE)
    method_heading = find_paragraph(doc, "METHODOLOGY")
    subsection_template = find_paragraph(doc, "Case-study model and lighting systems")
    body_template = find_paragraph(
        doc,
        "A three-dimensional model was constructed for a single open-plan office laboratory "
        "located within a larger building in Singapore, using the existing architectural geometry. "
        "Figure 1 identifies the simulated room within the host-building model; only this room, "
        "rather than the whole building, was included in the lighting simulations. The room contains "
        "24 fixed workstations, each equipped with an independently adjustable task light. The "
        "background system is divided into three ceiling-light zones: seats 0–7, 8–15, and 16–23. "
        "ClimateStudio was used to calculate illuminance and uniformity for the occupied task areas "
        "and their surroundings. Material reflectance, luminaire photometry, model orientation, and "
        "calculation-grid settings should be reported in the camera-ready version.",
    )
    caption_template = find_paragraph(
        doc, "Figure 1. Location of the simulated office room within the host building"
    )
    image_template = doc.paragraphs[24]

    renumber_existing_figures(doc)

    heading = new_paragraph_after(method_heading, subsection_template, OVERVIEW_HEADING)
    heading.paragraph_format.keep_with_next = True
    p1 = new_paragraph_after(heading, body_template, OVERVIEW_1)
    p2 = new_paragraph_after(p1, body_template, OVERVIEW_2)
    figure = add_method_figure(p2, image_template)
    add_caption(figure, caption_template)

    doc.save(OUTPUT)
    with zipfile.ZipFile(OUTPUT) as archive:
        bad_member = archive.testzip()
        if bad_member:
            raise ValueError(f"Corrupt ZIP member: {bad_member}")

    check = Document(OUTPUT)
    full_text = "\n".join(p.text for p in check.paragraphs)
    required = [OVERVIEW_HEADING, "Figure 1(a)", "Figure 1(f)", CAPTION_TEXT, "Figure 6."]
    for item in required:
        if item not in full_text:
            raise ValueError(f"Missing expected manuscript content: {item}")
    if len(check.inline_shapes) < 1:
        raise ValueError("The method diagram was not embedded")

    print(f"source={SOURCE}")
    print(f"output={OUTPUT}")
    print(f"paragraphs={len(check.paragraphs)} tables={len(check.tables)} inline_shapes={len(check.inline_shapes)}")


if __name__ == "__main__":
    main()
