from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import os
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent
TARGET = ROOT / "FullPaperTemplateASIM2026 - KK.docx"
TEMP = ROOT / "FullPaperTemplateASIM2026 - KK.revision_tmp.docx"
FIG_POWER = ROOT / "ASIM2026_figures" / "results_power_by_occupancy.png"
FIG_ALLOC = ROOT / "ASIM2026_figures" / "results_allocation_and_savings.png"


ABSTRACT = (
    "Traditional lighting design commonly assumes full occupancy, whereas workplace occupancy "
    "varies in both density and spatial distribution. This mismatch can cause unnecessary "
    "illumination in underused spaces. This study proposes a Balanced Light Control Framework "
    "that coordinates three ceiling-light zones and 24 workstation task lights while maintaining "
    "visual-comfort requirements. A three-dimensional model of a Singapore office laboratory was "
    "evaluated under Clear, Overcast, and Night conditions and occupancy counts from one to 15. "
    "Five matched seating distributions were generated at each count, producing 225 scenarios. "
    "A penalty-based genetic algorithm was compared with an idealized zonal PIR ceiling-only "
    "baseline and an occupancy-triggered rule-based hybrid baseline. Scenario-mean power for the "
    "GA-derived hybrid solutions was 17.44 W under Clear, 93.01 W under Overcast, and 236.10 W at "
    "Night. Relative to zonal PIR, these values were 76.9%, 62.0%, and 46.8% lower, respectively. "
    "Relative to rule-based hybrid control, the GA-derived solutions were 22.5% lower under "
    "Overcast and 25.0% lower at Night, but 6.7% higher under Clear. The results demonstrate a "
    "daylight-dependent transition from task-light-dominant to background-light-dominant operation, "
    "while the Clear-sky exception shows that optimization does not automatically outperform a "
    "well-specified fixed rule. Seat location also produced substantial within-count variation, "
    "supporting the use of spatial occupancy information in lighting control."
)


INTRO_CONTRIBUTION = (
    "This study asks: how should ceiling and task-light levels be coordinated under changing "
    "daylight and occupancy patterns to minimize lighting power while satisfying visual-comfort "
    "constraints? A Balanced Light Control Framework is evaluated in a simulated open-plan office. "
    "The contribution is threefold: (1) a joint representation of three ceiling-light zones and 24 "
    "task lights; (2) a penalty-based optimization formulation incorporating illuminance and "
    "uniformity; and (3) a matched comparison of GA-derived settings, zonal PIR ceiling-only control, "
    "and occupancy-triggered rule-based hybrid control across 225 lighting-occupancy-seating "
    "scenarios. The emphasis is a case-specific comparison rather than a deployed real-time controller."
)


METHOD_BASELINE_HEADING = "Rule-based comparison strategies"
METHOD_BASELINE_1 = (
    "Two deterministic controls were evaluated using the same 225 occupied-seat lists. In the "
    "idealized zonal PIR ceiling-only baseline, task lights remained off and an empty ceiling zone "
    "was switched off. For an occupied zone, the ceiling level was the maximum seat-specific "
    "requirement among its occupants. The occupancy-triggered rule-based hybrid used the same "
    "maximum rule for each zone's ceiling component, while every occupied workstation received its "
    "own prescribed task-light level and every unoccupied task light remained off. Both comparisons "
    "therefore use known seat locations; the hybrid strategy assumes workstation-level occupancy "
    "triggering rather than a conventional zone-only PIR sensor."
)
METHOD_BASELINE_2 = (
    "For zonal PIR under Night, the two outer seat columns in every four-seat row required ceiling "
    "level 13 and the two middle columns level 11. Under Overcast, seat-level requirements were "
    "0/3:10, 1/2:9, 4-6:8, 7:9, 8/11:8, 9/10:7, 12/15/16/19:5, "
    "13/14/17/18:4, and 20-23:0. Under Clear, they were 0-3:6, 4-5:1, "
    "6-7:2, 8-10:1, 11:2, and 12-23:0. For rule-based hybrid control, each pair "
    "(ceiling, task) was (8,3) at every occupied seat under Night. Under Overcast the pairs were "
    "0-3:(5,4), 4-7:(4,4), 8-11:(3,3), 12-19:(0,3), and 20-23:(0,0); "
    "under Clear they were 0-3:(0,4), 4-11:(0,1), 12-15:(0,0), "
    "16-19:(0,2), and 20-23:(0,0). "
    "These fixed levels were established for the stated visual requirements. Baseline power was "
    "then calculated with the same 13.2 W ceiling-zone and 1.5 W task-light level coefficients."
)


RESULTS_BLOCK = [
    ("heading", "Comparative power demand across strategies"),
    ("body", (
        "Figure 4 compares the three strategies as occupancy increased. Each line is the mean of five "
        "matched seat distributions and each band is their observed minimum-maximum range. Across all "
        "75 scenarios per condition, mean power for zonal PIR, rule-based hybrid, and GA-derived hybrid "
        "control was 75.50, 16.34, and 17.44 W under Clear; 244.46, 120.03, and 93.01 W under "
        "Overcast; and 443.52, 314.78, and 236.10 W at Night. Thus, GA-derived control reduced "
        "scenario-mean power relative to zonal PIR by 76.9%, 62.0%, and 46.8% for Clear, Overcast, "
        "and Night, respectively. Relative to rule-based hybrid control, it reduced mean power by "
        "22.5% under Overcast and 25.0% at Night, but increased it by 6.7% under Clear."
    )),
    ("figure4", None),
    ("caption4", None),
    ("heading", "Power allocation and paired strategy outcomes"),
    ("body", (
        "The daylight-dependent difference was primarily associated with ceiling-light demand "
        "(Figure 5a). The Clear GA-derived and rule-based hybrid solutions used no ceiling lighting "
        "on average and required 17.44 and 16.34 W of task lighting, respectively. Under Overcast, "
        "GA-derived ceiling power was 61.07 W compared with 86.77 W for rule-based hybrid control; "
        "at Night the corresponding values were 194.48 and 278.78 W. Task-light power was similar "
        "under Overcast (31.94 versus 33.26 W) and was higher for the GA-derived strategy at Night "
        "(41.62 versus 36.00 W), but the ceiling reduction dominated total power. Ceiling lighting "
        "therefore accounted for 0.0%, 65.7%, and 82.4% of GA-derived power under Clear, Overcast, "
        "and Night, respectively."
    )),
    ("body", (
        "Paired case comparisons reinforce the condition dependence. The GA-derived solution used "
        "less power than rule-based hybrid control in 7 of 75 Clear cases, 60 of 75 Overcast cases, "
        "and 69 of 75 Night cases; it used more power in 40, 3, and 0 cases, respectively, with the "
        "remaining 28, 12, and 6 cases tied. The negative mean Clear saving in Figure 5b therefore "
        "reflects a modest but recurring difference rather than a small number of extreme cases."
    )),
    ("figure5", None),
    ("caption5", None),
    ("heading", "Occupancy and seating-pattern sensitivity"),
    ("body", (
        "GA-derived mean power increased overall from one to 15 occupants, from 3.0 to 31.2 W under "
        "Clear, 38.8 to 125.1 W under Overcast, and 110.1 to 300.4 W at Night. The trajectories were "
        "not strictly monotonic because each occupancy count used different spatial layouts and "
        "activated different combinations of ceiling zones. Across occupancy counts, the mean "
        "within-count power range for the five layouts was 14.6 W under Clear, 24.8 W under "
        "Overcast, and 31.4 W at Night for the GA-derived strategy. The largest observed ranges were "
        "21.0 W at ten occupants under Clear, 76.5 W at two occupants under Overcast, and 85.2 W at "
        "three occupants at Night. Seat location therefore remained influential even when occupant "
        "count and sky condition were fixed."
    )),
]


DISCUSSION_PARAGRAPHS = [
    (
        "The comparison shows that the value of coordinated task and background lighting depends on "
        "daylight availability and on the reference controller. Relative to ceiling-only zonal PIR, "
        "the GA-derived strategy lowered scenario-mean power in all three conditions, with the largest "
        "relative reduction under Clear. Relative to the stronger rule-based hybrid baseline, however, "
        "additional optimization was beneficial only under Overcast and Night. In those conditions, "
        "the main mechanism was a lower ceiling-light contribution while task lights maintained local "
        "workplane requirements."
    ),
    (
        "The Clear result is an important counterexample to an assumption that a genetic algorithm "
        "must outperform a fixed rule. The rule-based hybrid required only localized task lighting and "
        "was lower than the retained GA solution in 40 of 75 cases, with 28 ties. If both strategies satisfy identical "
        "illuminance and uniformity constraints, this difference indicates that the present Galapagos "
        "runs should not be interpreted as strict global optima. Convergence, penalty weighting, discrete "
        "search resolution, and repeated-run variability should be examined. Accordingly, the results "
        "are described as GA-derived feasible solutions rather than a proven global performance bound."
    ),
    (
        "The variation among seating patterns also has a direct sensing implication. A controller based "
        "only on total occupancy would assign identical settings to spatially different cases, although "
        "zone activation and daylight exposure can produce markedly different power demand. Practical "
        "implementation may therefore benefit from zonal location or workstation-level occupancy "
        "information. The idealized baselines used known occupied-seat IDs; real PIR sensors with only "
        "zone-level detection would provide less spatial information and may require more conservative "
        "settings."
    ),
    (
        "Several limitations bound interpretation. The study represents one office geometry and three "
        "selected lighting conditions rather than annual daylight variation. Occupancy was limited to "
        "15 of 24 workstations, with five random layouts per count; these layouts are design samples, so "
        "their min-max ranges are not statistical confidence intervals. The comfort model considered "
        "simulated illuminance and uniformity but not individual preference, glare, sensor uncertainty, "
        "control delay, or hardware dynamics. Scenario means weight every count and layout equally and "
        "should not be interpreted as annual energy use. Finally, repeated GA runs and an independent "
        "cross-check of comfort compliance for all retained baseline and optimized settings are needed "
        "for a reproducibility audit."
    ),
]


CONCLUSION = (
    "This study evaluated coordinated control of three ceiling-light zones and 24 task lights across "
    "225 matched lighting-occupancy-seating scenarios. Compared with idealized zonal PIR ceiling-only "
    "control, the GA-derived hybrid strategy reduced scenario-mean power by 76.9% under Clear, 62.0% "
    "under Overcast, and 46.8% at Night. Compared with occupancy-triggered rule-based hybrid control, "
    "it reduced mean power by 22.5% under Overcast and 25.0% at Night, primarily through lower ceiling "
    "demand, but used 6.7% more power under Clear. This exception demonstrates that a simple "
    "daylight-specific rule can outperform the retained GA solution and that optimization claims must "
    "be supported by convergence and feasibility checks. The substantial variation among layouts at "
    "the same occupancy count further supports spatially resolved occupancy sensing. Overall, the "
    "framework provides a reproducible case-study platform for comparing occupancy-centric lighting "
    "strategies, while annual simulations and implementation-level validation remain necessary before "
    "claiming operational energy savings."
)


CAPTION4_LABEL = "Figure 4. "
CAPTION4_TEXT = (
    "Lighting power versus occupancy for zonal PIR ceiling-only, rule-based hybrid, and GA-derived "
    "hybrid control under (a) Clear, (b) Overcast, and (c) Night conditions. Lines show the mean "
    "across five matched seat distributions at each occupancy count; shaded bands show the observed "
    "minimum-maximum range (n = 5 layouts). Panel y-axis ranges differ to preserve visibility."
)
CAPTION5_LABEL = "Figure 5. "
CAPTION5_TEXT = (
    "Strategy-level power allocation and comparative savings. (a) Scenario-mean ceiling and task-light "
    "power, averaged equally across 75 occupancy-layout cases per condition. (b) Percentage reduction "
    "in GA-derived hybrid mean power relative to the two rule-based baselines; negative values denote "
    "higher GA-derived power."
)


def find_paragraph(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def clear_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def copy_run_format(source_run, target_run) -> None:
    if source_run is not None and source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def set_plain_text(paragraph: Paragraph, text: str, run_template=None) -> None:
    if run_template is None and paragraph.runs:
        run_template = paragraph.runs[0]
    clear_content(paragraph)
    run = paragraph.add_run(text)
    copy_run_format(run_template, run)


def new_paragraph_after(anchor: Paragraph, template: Paragraph, text: str = "") -> Paragraph:
    element = deepcopy(template._p)
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    set_plain_text(paragraph, text, template.runs[0] if template.runs else None)
    return paragraph


def new_caption_after(anchor: Paragraph, template: Paragraph, label: str, text: str) -> Paragraph:
    paragraph = new_paragraph_after(anchor, template, "")
    clear_content(paragraph)
    label_run = paragraph.add_run(label)
    label_run.bold = True
    label_run.italic = True
    text_run = paragraph.add_run(text)
    text_run.italic = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    return paragraph


def new_figure_after(anchor: Paragraph, template: Paragraph, image_path: Path,
                     alt_text: str, title: str) -> Paragraph:
    paragraph = new_paragraph_after(anchor, template, "")
    clear_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.85))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", title)
    return paragraph


def remove_between(start: Paragraph, end: Paragraph) -> None:
    node = start._p.getnext()
    while node is not None and node is not end._p:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node


def replace_section_body(doc: Document, heading_text: str, next_heading_text: str,
                         paragraphs: list[str], body_template: Paragraph) -> None:
    heading = find_paragraph(doc, heading_text)
    next_heading = find_paragraph(doc, next_heading_text)
    remove_between(heading, next_heading)
    anchor = heading
    for text in paragraphs:
        anchor = new_paragraph_after(anchor, body_template, text)


def main() -> None:
    if not FIG_POWER.exists() or not FIG_ALLOC.exists():
        raise FileNotFoundError("Required result figures are missing")
    doc = Document(TARGET)

    body_template = find_paragraph(doc, "This expression reproduced the reported power for all 225 exported solutions without error. Because the simulations represent discrete conditions without an assigned duration, the analysis reports electrical power (W), not energy consumption (kWh).")
    subsection_template = find_paragraph(doc, "Decision variables and lighting power")
    figure_template = doc.paragraphs[50]
    caption_template = find_paragraph(doc, "Figure 4. Optimized lighting power: (a) mean and full range across five seating distributions at each occupancy count; (b) mean allocation to ceiling and task lighting across all 75 scenarios per condition")

    # Abstract and introduction consistency.
    set_plain_text(doc.paragraphs[10], ABSTRACT)
    set_plain_text(doc.paragraphs[17], INTRO_CONTRIBUTION)
    set_plain_text(doc.paragraphs[12], "KEYWORDS Occupancy-centric lighting, task lighting, lighting optimization, PIR control, daylight availability")

    # Insert the reproducible baseline definitions immediately after the power model.
    insertion_anchor = body_template
    h = new_paragraph_after(insertion_anchor, subsection_template, METHOD_BASELINE_HEADING)
    p1 = new_paragraph_after(h, body_template, METHOD_BASELINE_1)
    new_paragraph_after(p1, body_template, METHOD_BASELINE_2)

    # Replace the complete Results section while retaining the surrounding manuscript structure.
    results_heading = find_paragraph(doc, "RESULTS")
    discussion_heading = find_paragraph(doc, "DISCUSSION")
    remove_between(results_heading, discussion_heading)
    anchor = results_heading
    for kind, text in RESULTS_BLOCK:
        if kind == "heading":
            anchor = new_paragraph_after(anchor, subsection_template, text or "")
        elif kind == "body":
            anchor = new_paragraph_after(anchor, body_template, text or "")
        elif kind == "figure4":
            anchor = new_figure_after(
                anchor, figure_template, FIG_POWER,
                "Three panels show mean lighting power versus occupancy under Clear, Overcast, and Night. "
                "GA-derived hybrid control is below zonal PIR on average, while rule-based hybrid is below "
                "the GA-derived strategy under Clear.",
                "Power by occupancy and control strategy",
            )
        elif kind == "caption4":
            anchor = new_caption_after(anchor, caption_template, CAPTION4_LABEL, CAPTION4_TEXT)
        elif kind == "figure5":
            anchor = new_figure_after(
                anchor, figure_template, FIG_ALLOC,
                "Stacked bars compare ceiling and task-light power; a second panel shows GA-derived power "
                "reductions relative to zonal PIR and rule-based hybrid control. The Clear comparison with "
                "rule-based hybrid is negative.",
                "Power allocation and baseline savings",
            )
        elif kind == "caption5":
            anchor = new_caption_after(anchor, caption_template, CAPTION5_LABEL, CAPTION5_TEXT)
        else:
            raise ValueError(kind)

    # Replace Discussion and Conclusion with baseline-aware interpretations.
    replace_section_body(doc, "DISCUSSION", "CONCLUSION AND IMPLICATIONS",
                         DISCUSSION_PARAGRAPHS, body_template)
    conclusion_heading = find_paragraph(doc, "CONCLUSION AND IMPLICATIONS")
    acknowledgements_heading = find_paragraph(doc, "ACKNOWLEDGEMENTS")
    remove_between(conclusion_heading, acknowledgements_heading)
    new_paragraph_after(conclusion_heading, body_template, CONCLUSION)

    # Keep caption/image pairs together and avoid isolated headings at page bottoms.
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in {"ABSTRACT", "INTRODUCTION", "METHODOLOGY", "RESULTS", "DISCUSSION",
                    "CONCLUSION AND IMPLICATIONS", "ACKNOWLEDGEMENTS", "REFERENCES",
                    METHOD_BASELINE_HEADING} or text in {
                    "Case-study model and lighting systems", "Scenario matrix",
                    "Decision variables and lighting power", "Visual-comfort constraints and objective function",
                    "Genetic-algorithm optimization", "Comparative power demand across strategies",
                    "Power allocation and paired strategy outcomes", "Occupancy and seating-pattern sensitivity",
                }:
            paragraph.paragraph_format.keep_with_next = True
        if text.startswith("Figure 4.") or text.startswith("Figure 5."):
            paragraph.paragraph_format.keep_together = True

    doc.save(TEMP)
    with zipfile.ZipFile(TEMP) as archive:
        bad = archive.testzip()
        if bad:
            raise ValueError(f"Corrupt ZIP member: {bad}")
    check = Document(TEMP)
    required_text = ["Figure 5.", "76.9%", "44.6%", METHOD_BASELINE_HEADING]
    full_text = "\n".join(p.text for p in check.paragraphs)
    for item in required_text:
        if item not in full_text:
            raise ValueError(f"Missing required manuscript content: {item}")
    if len(check.inline_shapes) < 3:
        raise ValueError("Expected building figure plus two new inline results figures")
    os.replace(TEMP, TARGET)
    print(f"updated={TARGET}")
    print(f"paragraphs={len(check.paragraphs)} tables={len(check.tables)} inline_shapes={len(check.inline_shapes)}")


if __name__ == "__main__":
    main()
