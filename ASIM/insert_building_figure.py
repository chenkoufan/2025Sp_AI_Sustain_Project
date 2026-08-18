from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
DOCX = ROOT / "FullPaperTemplateASIM2026 - KK.docx"
BACKUP = ROOT / "FullPaperTemplateASIM2026 - KK.before_building_insert.docx"
IMAGE = ROOT / "ASIM2026_figures" / "building.png"


def format_run(run, *, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic


if not BACKUP.exists():
    shutil.copy2(DOCX, BACKUP)

doc = Document(DOCX)

# Update the first methodological description without rebuilding the manuscript.
for p in doc.paragraphs:
    if p.text.startswith("A three-dimensional model was constructed for a single open-plan office laboratory"):
        revised = p.text.replace(
            "Figure 1a identifies the simulated room within the host-building model",
            "Figure 1 identifies the simulated room within the host-building model",
        )
        if revised != p.text:
            p.clear()
            format_run(p.add_run(revised))
        break

# The current room-model display follows the case-study paragraph as the first table.
target_table = doc.tables[0]._tbl

image_p = doc.add_paragraph()
image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
image_p.paragraph_format.space_before = Pt(3)
image_p.paragraph_format.space_after = Pt(0)
image_p.paragraph_format.keep_with_next = True
image_p.add_run().add_picture(str(IMAGE), width=Cm(12.0))

caption_p = doc.add_paragraph()
caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption_p.paragraph_format.line_spacing = 1.0
caption_p.paragraph_format.space_before = Pt(3)
caption_p.paragraph_format.space_after = Pt(6)
caption_run = caption_p.add_run(
    "Figure 1. Location of the simulated office room within the host building."
)
format_run(caption_run, bold=True, italic=True)

# Move the new paragraphs from the end of the document to immediately before the room-model table.
target_table.addprevious(image_p._p)
target_table.addprevious(caption_p._p)

# Renumber and correct the caption of the existing room-scale model.
for p in doc.paragraphs:
    if p.text.startswith("Figure 2. Case-study context and model:"):
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        format_run(
            p.add_run(
                "Figure 2. Room-scale simulation model, workstation layout, and lighting-system arrangement."
            ),
            bold=True,
            italic=True,
        )
        break

doc.save(DOCX)
