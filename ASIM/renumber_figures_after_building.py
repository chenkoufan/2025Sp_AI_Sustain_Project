from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path(__file__).resolve().parent / "FullPaperTemplateASIM2026 - KK.docx"
doc = Document(DOCX)


def set_font(run, *, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


for p in doc.paragraphs:
    text = p.text
    revised = text
    revised = revised.replace(
        "Figure 2 documents a representative seven-occupant case",
        "Figure 3 documents a representative seven-occupant case",
    )
    revised = revised.replace("Figure 3a", "Figure 4a")
    revised = revised.replace("Figure 3b", "Figure 4b")

    if text.startswith("Figure 2. Representative seven-occupant optimization:"):
        revised = text.replace("Figure 2.", "Figure 3.", 1)
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(revised), bold=True, italic=True)
    elif text.startswith("Figure 3. Optimized lighting power:"):
        revised = text.replace("Figure 3.", "Figure 4.", 1)
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        set_font(p.add_run(revised), bold=True, italic=True)
    elif revised != text:
        p.clear()
        set_font(p.add_run(revised))

doc.save(DOCX)
