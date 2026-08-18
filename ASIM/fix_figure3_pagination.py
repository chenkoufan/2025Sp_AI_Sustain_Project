from __future__ import annotations

from pathlib import Path
import os
import zipfile

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
TARGET = ROOT / "FullPaperTemplateASIM2026 - KK.docx"
TEMP = ROOT / "FullPaperTemplateASIM2026 - KK.pagination_tmp.docx"
CAPTION = (
    "Figure 3. Representative seven-occupant optimization: (a) complete Galapagos interface; "
    "(b) corresponding occupant distribution and simulated illuminance field"
)


def main() -> None:
    doc = Document(TARGET)
    body = doc._element.body
    caption_element = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == CAPTION:
            caption_element = paragraph._p
            break
    if caption_element is None:
        raise ValueError("Figure 3 caption not found")

    spacer = caption_element.getnext()
    if spacer is None or spacer.tag != qn("w:p"):
        raise ValueError("Expected spacer paragraph after Figure 3 caption")
    p_pr = spacer.find(qn("w:pPr"))
    if p_pr is not None:
        keep_next = p_pr.find(qn("w:keepNext"))
        if keep_next is not None:
            p_pr.remove(keep_next)

    doc.save(TEMP)
    with zipfile.ZipFile(TEMP) as archive:
        bad = archive.testzip()
        if bad:
            raise ValueError(f"Corrupt ZIP member: {bad}")
    Document(TEMP)
    os.replace(TEMP, TARGET)
    print("removed_figure3_results_keep_chain=true")


if __name__ == "__main__":
    main()
